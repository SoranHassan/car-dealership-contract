# word/generator.py
#
# Cross-platform contract generator (Linux + Windows + macOS).
#
# The previous implementation drove Microsoft Word through COM automation
# (pywin32 / win32com), which only works on Windows with Office installed.
# This version uses python-docx, so it runs anywhere Python does.
#
# The template (assets/AG.docx) stores its fields inside Word text boxes.
# python-docx does not expose text boxes through its high-level API, so we
# walk the raw OOXML tree (every <w:p> paragraph, including those nested in
# <w:txbxContent>). Placeholders are often split across several <w:t> runs,
# so we coalesce each paragraph's runs before replacing.

import os
import json
import shutil
import tempfile
import logging
import sys

from persiantools.jdatetime import JalaliDate

from docx import Document
from docx.shared import Cm, Emu
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


def resource_path(relative_path):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(relative_path)


class ContractGeneratorError(Exception):
    pass


class ContractGenerator:
    # Fallback size for the inspection photo, only used if the enclosing
    # text box's real dimensions can't be read from the template (should
    # not normally happen — see _box_extent_emu).
    CHECKPOINT_IMG_WIDTH_CM = 6.5

    def __init__(self):
        self.word_template = resource_path("./assets/AG.docx")
        self._check_template_exists()

    def _check_template_exists(self):
        if not os.path.exists(self.word_template):
            raise ContractGeneratorError(f"Template not found: {self.word_template}")

    # ------------------------------------------------------------------ #
    # Public API                                                          #
    # ------------------------------------------------------------------ #
    def generate(self, json_path, checkpoint_image_path, output_dir):
        temp_dir = None
        try:
            import time
            start_time = time.time()

            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            flat = self.flatten_data(data)

            buyer_ncode = flat.get("buyer_ncode", "unknown")
            seller_ncode = flat.get("seller_ncode", "unknown")
            contract_number = flat.get("deal_num", "unknown")
            file_name = f"{buyer_ncode} - {seller_ncode} - {contract_number}.docx"

            date_folder = flat.get("deal_date", "unknown").replace("/", "-")
            final_output_dir = os.path.join(output_dir, date_folder)
            os.makedirs(final_output_dir, exist_ok=True)
            final_output = os.path.join(final_output_dir, file_name)

            temp_dir = tempfile.mkdtemp(prefix="contract_gen_")
            temp_output = os.path.join(temp_dir, "temp_output.docx")

            doc = Document(os.path.abspath(self.word_template))
            self._checkpoint_image_inserted = False
            self._fill_document(doc, flat, checkpoint_image_path)
            doc.save(temp_output)

            if os.path.exists(final_output):
                try:
                    os.remove(final_output)
                except OSError:
                    pass

            shutil.copy2(temp_output, final_output)

            logger.info(
                f"Contract generated in {time.time() - start_time:.2f}s: {final_output}"
            )
            return final_output

        except ContractGeneratorError:
            raise
        except Exception as e:
            logger.error(f"Error in generate: {e}")
            raise ContractGeneratorError(f"خطا در تولید قرارداد: {str(e)}")
        finally:
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except OSError:
                pass

    # ------------------------------------------------------------------ #
    # Template filling                                                    #
    # ------------------------------------------------------------------ #
    def _fill_document(self, doc, flat, checkpoint_image_path):
        """Replace every placeholder across the whole document, including
        paragraphs nested inside text boxes."""
        # Replace longest keys first so a shorter key can never clobber part
        # of a longer one (e.g. seller_from vs seller_fname).
        ordered_keys = sorted(flat.keys(), key=len, reverse=True)

        body = doc.element.body
        for p in body.iter(qn("w:p")):
            self._process_paragraph(p, flat, ordered_keys, checkpoint_image_path, doc)

    @staticmethod
    def _own_text_elems(p):
        """Return the <w:t> elements that belong directly to this paragraph's
        own runs (and hyperlink runs), excluding any inside nested text-box
        paragraphs. Using descendant './/w:t' would double-count text boxes,
        because their inner <w:p> is visited separately."""
        elems = []
        for child in p:
            if child.tag == qn("w:r"):
                elems.extend(child.findall(qn("w:t")))
            elif child.tag == qn("w:hyperlink"):
                for r in child.findall(qn("w:r")):
                    elems.extend(r.findall(qn("w:t")))
        return elems

    def _process_paragraph(self, p, flat, ordered_keys, checkpoint_image_path, doc):
        t_elems = self._own_text_elems(p)
        if not t_elems:
            return

        original = "".join(t.text or "" for t in t_elems)
        if not original:
            return

        new_text = original

        # 1) Inspection photo: drop the tag and insert the image in this box.
        # The template stores each shape twice for compatibility (a modern
        # DrawingML "Choice" and a legacy VML "Fallback" of the same visual
        # shape via mc:AlternateContent) — both contain the placeholder text,
        # but only the Choice (processed first, in document order) is ever
        # actually rendered by Word/LibreOffice, so the picture is inserted
        # only once, into that occurrence.
        if "checkpoint_img" in new_text:
            new_text = new_text.replace("checkpoint_img", "")
            if (
                not self._checkpoint_image_inserted
                and checkpoint_image_path
                and os.path.exists(checkpoint_image_path)
            ):
                try:
                    self._insert_image(p, checkpoint_image_path, doc)
                    self._checkpoint_image_inserted = True
                except Exception as img_error:  # pragma: no cover - defensive
                    logger.warning(f"Failed to place image: {img_error!r}")

        # 2) Payment stamp.
        if "paid_stamp" in new_text:
            status = "پرداخت شد" if flat.get("is_payed") in [1, "1", True] else "پرداخت نشد"
            new_text = new_text.replace("paid_stamp", status)

        # 3) All remaining field placeholders.
        for key in ordered_keys:
            if key in new_text:
                new_text = new_text.replace(key, str(flat[key]))

        if new_text != original:
            self._set_paragraph_text(t_elems, new_text)

    @staticmethod
    def _set_paragraph_text(t_elems, new_text):
        """Write the coalesced text back into the first run and clear the rest,
        preserving whitespace."""
        first = t_elems[0]
        first.text = new_text
        first.set(qn("xml:space"), "preserve")
        for t in t_elems[1:]:
            t.text = ""

    def _box_extent_emu(self, p):
        """Walk up from a paragraph nested inside a text box to the enclosing
        <wp:inline>/<wp:anchor> and read its <wp:extent> — the real,
        designed size (in EMU) of the box the paragraph lives in. Returns
        (cx, cy) or None if the paragraph isn't inside a drawing (shouldn't
        happen for our template, but we don't want to ever raise here)."""
        node = p
        while node is not None:
            if node.tag in (qn("wp:inline"), qn("wp:anchor")):
                extent = node.find(qn("wp:extent"))
                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    if cx and cy:
                        return int(cx), int(cy)
                return None
            node = node.getparent()
        return None

    def _insert_image(self, p, image_path, doc):
        para = Paragraph(p, doc)
        run = para.add_run()
        extent = self._box_extent_emu(p)
        if extent:
            cx, cy = extent
            run.add_picture(image_path, width=Emu(cx), height=Emu(cy))
        else:
            logger.warning(
                "checkpoint_img box extent not found in template; "
                "falling back to a fixed width"
            )
            run.add_picture(image_path, width=Cm(self.CHECKPOINT_IMG_WIDTH_CM))

    # ------------------------------------------------------------------ #
    # Data flattening (unchanged public contract)                        #
    # ------------------------------------------------------------------ #
    def flatten_data(self, data):
        flat = {}

        # Seller
        flat["seller_birth"] = data["seller"].get("birth", "")
        flat["seller_fname"] = f"{data['seller'].get('name', '')} {data['seller'].get('lname', '')}"
        flat["seller_ncode"] = data["seller"].get("national_code", "")
        flat["seller_shcode"] = data["seller"].get("national_shcode", "")
        flat["seller_from"] = data["seller"].get("from", "")
        flat["seller_phone"] = data["seller"].get("phone", "")
        flat["seller_adress"] = data["seller"].get("adress", "")
        flat["seller_father"] = data["seller"].get("father", "")

        # Buyer
        flat["buyer_birth"] = data["buyer"].get("birth", "")
        flat["buyer_fname"] = f"{data['buyer'].get('name', '')} {data['buyer'].get('lname', '')}"
        flat["buyer_ncode"] = data["buyer"].get("national_code", "")
        flat["buyer_shcode"] = data["buyer"].get("national_shcode", "")
        flat["buyer_from"] = data["buyer"].get("from", "")
        flat["buyer_phone"] = data["buyer"].get("phone", "")
        flat["buyer_adress"] = data["buyer"].get("address", "")
        flat["buyer_father"] = data["buyer"].get("father", "")

        # Car
        flat["car_type"] = data["car_deal"].get("type", "")
        flat["car_color"] = data["car_deal"].get("color", "")
        flat["car_system"] = data["car_deal"].get("system", "")
        flat["car_model"] = data["car_deal"].get("model", "")
        flat["body_id"] = data["car_deal"].get("body_id", "")
        flat["motor_id"] = data["car_deal"].get("motor_id", "")
        flat["car_kilometer"] = data["car_deal"].get("kilometer", "")
        flat["pelak"] = data["car_deal"].get("pelak", "")
        flat["car_info"] = data["car_deal"].get("car_info", "")

        # Deal Info
        flat["deal_time"] = data["deal_info"].get("deal_time", "")
        flat["day_respite"] = data["deal_info"].get("day_respite", "")
        flat["price_rial"] = data["deal_info"].get("price_rial", "")
        flat["price_toman"] = data["deal_info"].get("price_toman", "")
        flat["price_info"] = data["deal_info"].get("price_info", "")
        flat["description_text"] = data["deal_info"].get("description_text", "")
        flat["deal_date"] = JalaliDate.today().strftime("%Y/%m/%d")
        flat["deal_num"] = data["deal_info"].get("deal_num", "")

        is_payed = data["deal_info"].get("is_payed")
        flat["is_payed"] = 1 if is_payed in [1, "1", True] else 0

        return flat


def generate_contract(json_path, checkpoint_image_path, output_dir):
    """تابع کمکی برای تولید قرارداد"""
    generator = ContractGenerator()
    return generator.generate(json_path, checkpoint_image_path, output_dir)
